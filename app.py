#-*- coding: utf-8 -*-
import random,os,json,io,re,zipfile,tempfile
import ssl
import pandas as pd
import streamlit as st
import streamlit_toggle as tog
from langchain import HuggingFaceHub
from langchain.llms import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from utils import text_to_docs
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from io import StringIO
from io import BytesIO
from usellm import Message, Options, UseLLM
from huggingface_hub import login
# import cv2
# import pdfplumber
# import pytesseract
# from pdf2image import convert_from_path
#from playsound import playsound
#from langchain.text_splitter import CharacterTextSplitter
#from langchain.embeddings.openai import OpenAIEmbeddings
#from langchain.chains.summarize import load_summarize_chain
#import os
#import pyaudiogi
#import wave
#from langchain.document_loaders import UnstructuredPDFLoader
#import streamlit.components.v1 as components
#from st_custom_components import st_audiorec, text_to_docs
#import sounddevice as sd
#from scipy.io.wavfile import write

# Setting Env
if st.secrets["OPENAI_API_KEY"] is not None:
    os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
else:
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPENAI_API_KEY")

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    return output_pdf


@st.cache_data
def usellm(prompt):
    """
    Getting GPT-3.5 Model into action
    """
    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a fraud analyst, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

# Setting Config for Llama-2
login(token=st.secrets["HUGGINGFACEHUB_API_TOKEN"])
os.environ["HUGGINGFACEHUB_API_TOKEN"] = st.secrets["HUGGINGFACEHUB_API_TOKEN"]


llama_13b = HuggingFaceHub(
            repo_id="meta-llama/Llama-2-13b-chat-hf",
            model_kwargs={"temperature":0.01, 
                        "min_new_tokens":100, 
                        "max_new_tokens":500})

memory = ConversationSummaryBufferMemory(llm= llama_13b, max_token_limit=500)
conversation = ConversationChain(llm= llama_13b, memory=memory,verbose=False)


@st.cache_data
def llama_llm(_llm,prompt):
    response = _llm.predict(prompt)
    return response

@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text

@st.cache_resource
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings

# @st.cache_data
# def embedding_store(pdf_files):
#     merged_pdf = merge_pdfs(pdf_files)
#     final_pdf = PyPDF2.PdfReader(merged_pdf)
#     text = ""
#     for page in final_pdf.pages:
#         text += page.extract_text()
#     texts =  text_splitter.split_text(text)
#     docs = text_to_docs(texts)
#     docsearch = FAISS.from_documents(docs, hf_embeddings)
#     return docs, docsearch


@st.cache_data
def embedding_store(pdf_files):
    pdf_only =[]
    for file in pdf_files:
      if file.endswith('.pdf'):
        pdf_only.append(file)       
      
    merged_pdf = merge_pdfs(pdf_only)
    final_pdf = PyPDF2.PdfReader(merged_pdf)
    text = ""
    for page in final_pdf.pages:
        text += page.extract_text()
      
    for file in pdf_files:
      if file.endswith('xlsx'):
        df = pd.read_excel(file, engine='openpyxl')
        # Find the row index where the table data starts
        data_start_row = 0  # Initialize to 0
        for i, row in df.iterrows():
            if row.notna().all():
                data_start_row = i
                break
              
        if data_start_row>0:  
            df.columns = df.iloc[data_start_row]
          
        
        # Extract the text content above the data
        text += "\n".join(df.iloc[:data_start_row].apply(lambda x: "\t".join(map(str, x)), axis=1)).replace('nan','')
        
        df = df.iloc[data_start_row+1:]
        text_buffer = StringIO()
        df.to_csv(text_buffer, sep='\t', index=False)
        text += "\n\n"+ text_buffer.getvalue()
        text_buffer.close()
        
    texts =  text_splitter.split_text(text)
    docs = text_to_docs(texts)
    docsearch = FAISS.from_documents(docs, hf_embeddings)
    return docs, docsearch
    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()

# def merge_and_extract_text(pdf_list):
#     merged_pdf = fitz.open()
#     # Merge the PDF files
#     for pdf_file in pdf_list:
#         pdf_document = fitz.open(pdf_file)
#         merged_pdf.insert_pdf(pdf_document)
#     # Create an empty string to store the extracted text
#     merged_text = ""
#     # Extract text from each page of the merged PDF
#     for page_num in range(merged_pdf.page_count):
#         page = merged_pdf[page_num]
#         text = page.get_text()
#         merged_text += text
#     # Close the merged PDF
#     merged_pdf.close()
#     return merged_text


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False



def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            all_text.append(text)
    return "\n".join(all_text)




# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# convert scanned pdf to searchable pdf
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    print("Running OCR")
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text


# Setting globals
if "visibility" not in st.session_state:
    st.session_state.visibility = "visible"
    st.session_state.disabled = True
if "stored_session" not in st.session_state:
    st.session_state["stored_session"] = []

if "tmp_table_gpt_fd" not in st.session_state:
    st.session_state.tmp_table_gpt_fd = pd.DataFrame()
if "tmp_table_llama_fd" not in st.session_state:
    st.session_state.tmp_table_llama_fd = pd.DataFrame()
if "tmp_table_gpt_aml" not in st.session_state:
    st.session_state.tmp_table_gpt_aml = pd.DataFrame()
if "tmp_table_llama_aml" not in st.session_state:
    st.session_state.tmp_table_llama_aml = pd.DataFrame()

if "tmp_summary_gpt_fd" not in st.session_state:
    st.session_state["tmp_summary_gpt_fd"] = ''
if "tmp_summary_llama_fd" not in st.session_state:
    st.session_state["tmp_summary_llama_fd"] = ''
if "tmp_summary_gpt_aml" not in st.session_state:
    st.session_state["tmp_summary_gpt_aml"] = ''
if "tmp_summary_llama_aml" not in st.session_state:
    st.session_state["tmp_summary_llama_aml"] = ''

if "tmp_narrative_gpt" not in st.session_state:
    st.session_state["tmp_narrative_gpt"] = ''
if "tmp_narrative_llama" not in st.session_state:
    st.session_state["tmp_narrative_llama"] = ''

if "case_num" not in st.session_state:
    st.session_state.case_num = ''
if "fin_opt" not in st.session_state:
    st.session_state.fin_opt = ''
if "context_1" not in st.session_state:
    st.session_state.context_1 = ''
if "llm" not in st.session_state:
    st.session_state.llm = 'Open-AI'
if "pdf_files" not in st.session_state:
    st.session_state.pdf_files = []



# Apply CSS styling to resize the buttons
st.markdown("""
    <style>
        .stButton button {
            width: 145px;
            height: 35px;
        }
    </style>
""", unsafe_allow_html=True)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)

    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)

@st.cache_data
def create_zip_file(file_paths, zip_file_name):
    with zipfile.ZipFile(zip_file_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path, os.path.basename(file_path))



####### This markdown is to manage app style (app markdown)
st.markdown("""

<style>

.st-d5 {
    line-height: 1;
}


.css-1upf7v9 { 
    gap: 0.5rem;
}

.css-1balh2r{
    gap: 0;
}

.css-1544g2n {
    padding: 0;
    padding-top: 2rem;
    padding-right: 1rem;
    padding-bottom: 1.5rem;
    padding-left: 1rem;
}

.css-1q2g7hi {
    top: 2px;
    min-width: 350px;
    max-width: 600px;
    }

.st-ah {
    line-height: 1;
}

.st-af {
    font-size: 1.5rem;
}

.css-1a65djw {
    gap: 0;
    }

.css-1y4p8pa {
    width: 100%;
    padding: 3rem 1rem 10rem;
    padding-top: 3rem;
    padding-bottom: 10rem;
    max-width: 60rem;
}

.css-xujc5b p{
font-size: 25px;
}

.css-jzprzu {
    height: 2rem;
    min-height: 1.5rem;
    }

</style>
""", unsafe_allow_html=True)



# Addding markdown styles(Global)
st.markdown("""
<style>
.big-font {
    font-size:60px !important;
}
</style>
""", unsafe_allow_html=True)


# Set Sidebar
st.markdown("""
<style>
    [data-testid=stSidebar] {
        background-color: FBFBFB;
    }
</style>
""", unsafe_allow_html=True)

#Adding llm type-> st.session_state.llm
st.session_state.llm = st.radio("",options = pd.Series(["","Open-AI","Open-Source"]), horizontal=True)

st.markdown("---")

st.title("Suspicious Activity Reporting Assistant")
with st.sidebar:
    # st.sidebar.write("This is :blue[test]")
    # Navbar
    st.markdown('<link rel="stylesheet" href="https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css" integrity="sha384-Gn5384xqQ1aoWXA+058RXPxPg6fy4IWvTNh0E263XmFcJlSAwiGgFAW/dAiS6JXm" crossorigin="anonymous">', unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-top navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <style>
    .navbar-brand img {
      max-height: 50px; /* Adjust the height of the logo */
      width: auto; /* Allow the width to adjust based on the height */
    }
    </style>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
            <a class="navbar-brand" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
                <strong>| Operations Process Automation</strong>
            </a>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <nav class="navbar fixed-bottom navbar-expand-lg navbar-dark" style="background-color: #000000;">
    <button class="navbar-toggler" type="button" data-toggle="collapse" data-target="#navbarNav" aria-controls="navbarNav" aria-expanded="false" aria-label="Toggle navigation">
        <span class="navbar-toggler-icon"></span>
    </button>
    <div class="collapse navbar-collapse" id="navbarNav">
        <ul class="navbar-nav">
        <li class="nav-item active">
        <!--p style='color: white;'><b>Powered by EXL</b></p--!>
        <p style='color: white;'> <strong>Powered by EXL</strong> </p>
            <!--a class="nav-link disabled" href="#">
                <img src="https://www.exlservice.com/themes/exl_service/exl_logo_rgb_orange_pos_94.png" width="50" height="30" alt="">
                <span class="sr-only">(current)</span>
            </a--!>
        </li>
        </ul>
    </div>
    </nav>
    """, unsafe_allow_html=True)

    # Add the app name
    st.sidebar.markdown('<p class="big-font">SARA</p>', unsafe_allow_html=True)
    # st.sidebar.header("SARA")
    st.markdown("---")

    # Add a drop-down for case type
    options1 = ["Select Case Type", "Fraud transaction dispute", "AML"]
    selected_option_case_type = st.sidebar.selectbox("", options1)
    st.markdown("---")
    
    # Add a single dropdown
    options2 = ["Select Case ID", "SAR-2023-24680", "SAR-2023-13579", "SAR-2023-97531", "SAR-2023-86420", "SAR-2023-24681"]
    selected_option = st.sidebar.selectbox("", options2)
    # Add the image to the sidebar below options
    st.sidebar.image("MicrosoftTeams-image (3).png", use_column_width=True)

    
# Assing action to the main section
if selected_option_case_type == "Select Case Type":
    st.header("")

## Fraud Transaction Code started
elif selected_option_case_type == "Fraud transaction dispute":
    st.markdown("### :blue[Fraud transaction dispute]")

# st.markdown('---')

    # Redirect to Merge PDFs page when "Merge PDFs" is selected
    if selected_option == "SAR-2023-24680":
        st.session_state.case_num = "SAR-2023-24680"
        # st.header("Merge Documents")
        # st.write("Upload multiple document files and merge them into one doc.")

        # Upload PDF files
        # st.subheader("Upload Case Files")
        # st.markdown(f"**Case No: {st.session_state.case_num}**")
        # st.markdown("""
        #     | Case No.                  | Case Type                 | Customer Name             | Case Status             | Open Date              |
        #     | ------------------------  | ------------------------- | ------------------------- | ------------------------|------------------------|
        #     | SAR-2023-24680            | Fraud Transaction Dispute | John Brown                | In Progress             | 12/10/2020             |
        #     """)

        col1,col2 = st.columns(2)
        # Row 1
        with col1:
            st.markdown("##### **Case number&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** SAR-2023-24680")
            st.markdown("##### **Customer name  :** John Brown")


        with col2:
            st.markdown("##### **Case open date&nbsp;&nbsp;&nbsp;&nbsp;:** Feb 02, 2021")
            st.markdown("##### **Case type&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Fraud transaction")


        # Row 2
        with col1:
            st.markdown("##### **Customer ID&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** 9659754")


        with col2:
            st.markdown("##### **Case Status&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;:** Open")

        st.markdown("---")

        ## Defining some global variables for fraud transaction

        directoty_path = "data/"
        fetched_files = read_pdf_files(directoty_path)
        
        if selected_option:
            col1_up, col2_up, col3_up, col4_up, col5_up = st.tabs(["Data", "Generate Insights","Summarization","Download Report", "Make a Decision"])
            
            with col1_up:
                bt1_up, bt2_up = st.tabs(["Fetch Evidence", "Upload Evidence"])

                with bt1_up:
                    # Set the color
                    # st.markdown(
                    #     """
                    #     <div style="display: flex; justify-content: center; align-items: center; height: 48px; border: 1px solid #ccc; border-radius: 5px; background-color: #f2f2f2;">
                    #         <span style="font-size: 16px;  ">Fetch Evidence</span>
                    #     </div>
                    #     """,
                    #     unsafe_allow_html=True
                    # )
                    if 'clicked' not in st.session_state:
                        st.session_state.clicked = False
                    
                    def set_clicked():
                        st.session_state.clicked = True
                        st.session_state.disabled = True
                    st.write("") #for the gap
                    st.button('Fetch Evidence', on_click=set_clicked)

                    if st.session_state.clicked:
                        # st.write("Evidence Files:") 
                        # st.markdown(html_str, unsafe_allow_html=True)
                        
                        # Showing files
                        # show_files = fetched_files.copy()
                        # show_files = show_files + ['Other.pdf']
                        # files_frame = pd.DataFrame(show_files, columns=["File Name"])
                        # # files_frame["Select"] = [True for _ in range(len(files_frame))]
                        # files_frame = files_frame.reset_index(drop=True)

                        # # Add checkboxes to the DataFrame
                        # df_with_checkboxes = add_checkboxes_to_dataframe(files_frame)
                        
                        # # Iterate through each row and add checkboxes
                        # for index, row in df_with_checkboxes.iterrows():
                        #     if index < len(df_with_checkboxes) - 1:
                        #         checkbox_state = st.checkbox(f" {row['File Name']}", value=True)
                        #         df_with_checkboxes.loc[index, 'Select'] = checkbox_state
                        #     else:
                        #         st.checkbox(f"{row['File Name']}", value=False)



                        # st.dataframe(files_frame)
                        # st.write(df_reset.to_html(index=False), unsafe_allow_html=True)
                        # st.markdown(files_frame.style.hide(axis="index").to_html(), unsafe_allow_html=True)
                        
                        
                        
                        #select box to select file
                        selected_file_name = st.selectbox(":blue[Select a file to View]",fetched_files)
                        st.write("Selected File: ", selected_file_name)
                        st.session_state.disabled = False
                        file_ext = tuple("pdf")
                        if selected_file_name.endswith(file_ext):
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            #converting pdf data to bytes so that render_pdf_as_images could read it
                            file = pdf_to_bytes(selected_file_path)
                            pdf_images = render_pdf_as_images(file)
                            #showing content of the pdf
                            st.subheader(f"Contents of {selected_file_name}")
                            for img_bytes in pdf_images:
                                st.image(img_bytes, use_column_width=True)
                        else:
                            selected_file_path = os.path.join(directoty_path, selected_file_name)
                            # This is showing png,jpeg files
                            st.image(selected_file_path, use_column_width=True)



                with bt2_up:
                    pdf_files = st.file_uploader("", type=["pdf","png","jpeg","docx","xlsx"], accept_multiple_files=True)
                    st.session_state.pdf_files = pdf_files
                    # showing files
                    for uploaded_file in pdf_files:
                        #This code is to show pdf files
                        file_ext = tuple("pdf")
                        if uploaded_file.name.endswith(file_ext):
                            # Show uploaded files in a dropdown
                            # if pdf_files:
                            st.subheader("Uploaded Files")
                            file_names = [file.name for file in pdf_files]
                            selected_file = st.selectbox(":blue[Select a file]", file_names)
                            # Enabling the button
                            st.session_state.disabled = False
                            # Display selected PDF contents
                            if selected_file:
                                selected_pdf = [pdf for pdf in pdf_files if pdf.name == selected_file][0]
                                pdf_images = render_pdf_as_images(selected_pdf)
                                st.subheader(f"Contents of {selected_file}")
                                for img_bytes in pdf_images:
                                    st.image(img_bytes, use_column_width=True)

                        else:
                            # This is showing png,jpeg files
                            st.image(uploaded_file, use_column_width=True)

            #creating temp directory to have all the files at one place for accessing
                tmp_dir_ = tempfile.mkdtemp()
                temp_file_path= []


                for uploaded_file in pdf_files:
                    file_ext = tuple("pdf")
                    if uploaded_file.name.endswith(file_ext):
                        file_pth = os.path.join(tmp_dir_, uploaded_file.name)
                        with open(file_pth, "wb") as file_opn:
                            file_opn.write(uploaded_file.getbuffer())
                            temp_file_path.append(file_pth)
                    else:
                        pass

                for fetched_pdf in fetched_files:
                    file_ext = tuple("pdf")
                    if fetched_pdf.endswith(file_ext):
                        file_pth = os.path.join('data/', fetched_pdf)
                        # st.write(file_pth)
                        temp_file_path.append(file_pth) 
                    else:
                        pass   
                    
                #combining files in fetch evidence and upload evidence
                pdf_files_ = []
                if temp_file_path:
                    if pdf_files and fetched_files:
                        file_names = [file.name for file in pdf_files]
                        file_names = file_names + fetched_files
                        pdf_files_ = file_names
                    elif fetched_files:
                        pdf_files_ = fetched_files
                    elif pdf_files:
                        pdf_files_ = pdf_files
                    else: 
                        pass
            


            with col2_up:
                #This is the embedding model
                model_name = "sentence-transformers/all-MiniLM-L6-v2"
                # model_name = "hkunlp/instructor-large"
                
                # Memory setup for gpt-3.5
                llm = ChatOpenAI(temperature=0.1)
                memory = ConversationSummaryBufferMemory(llm=llm, max_token_limit=500)
                conversation = ConversationChain(llm=llm, memory =memory,verbose=False)
                
                
                # Adding condition on embedding
                try:
                    if temp_file_path:
                        hf_embeddings = embed(model_name) 
                    else:
                        pass
                except NameError:
                    pass
                
                # Chunking with overlap
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size = 1000,
                    chunk_overlap  = 100,
                    length_function = len,
                    separators=["\n\n", "\n", " ", ""]
                )
               

               # Creating header
                col1,col2 = st.columns(2)
                with col1:
                    st.markdown("""<span style="font-size: 24px; ">Pre-Set Questionnaire</span>""", unsafe_allow_html=True)
                    # Create a Pandas DataFrame with your data
                    data = {'Questions': [" What is the victim's name?","What is the suspect's name?",' List the merchant name',' How was the bank notified?',' When was the bank notified?',' What is the fraud type?',' When did the fraud occur?',' Was the disputed amount greater than 5000 USD?',' What type of cards are involved?',' Was the police report filed?']}
                    df_fixed = pd.DataFrame(data)
                    df_fixed.index = df_fixed.index +1
                with col2:
                    # Create a checkbox to show/hide the table
                    cols1, cols2, cols3, cols4 = st.columns([1,1,1,1])
                    with cols1:
                        show_table1 = tog.st_toggle_switch(label="", 
                                            key="Key1", 
                                            default_value=False, 
                                            label_after = False, 
                                            inactive_color = '#D3D3D3', 
                                            active_color="#11567f", 
                                            track_color="#29B5E8"
                                            )
                    # Show the table if the checkbox is ticked
                    if show_table1:
                        df_fixed["S.No."] = df_fixed.index
                        df_fixed = df_fixed.loc[:,['S.No.','Questions']]
                        st.markdown(df_fixed.style.hide(axis="index").to_html(), unsafe_allow_html=True)
            



